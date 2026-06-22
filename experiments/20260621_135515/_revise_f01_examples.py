from __future__ import annotations

import csv
import hashlib
from copy import copy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_BREAK
from docx.shared import Pt
from openpyxl import load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side


ROOT = Path(__file__).resolve().parents[2]
EXP = Path(__file__).resolve().parent
PACKET = EXP / "teacher_packet"
BENCHMARK = EXP / "benchmark"
DELIVERABLES = EXP / "deliverables"
MATERIAL_ROOT = ROOT / "document" / "teacher_training_curriculum" / "TapHuan-GV-TinHoc9"


FAILURE_CODES = [
    ("CF01_SAI_KIEN_THUC", "Sai kiến thức trọng yếu hoặc tạo hiểu sai mới."),
    ("CF02_VUOT_PHAM_VI_LOP_9", "Đòi hỏi kiến thức vượt rõ phạm vi lớp 9 mà không có lí do."),
    ("CF03_BO_QUA_DU_KIEN_HOC_SINH", "Bỏ qua bài làm, câu trả lời hoặc điểm vướng đã được cung cấp."),
    ("CF04_BIA_DU_KIEN_HOAC_KET_QUA", "Bịa dữ kiện, kết quả chạy hoặc phần sản phẩm chưa được quan sát."),
    ("CF05_TIET_LO_TOAN_BO_LOI_GIAI", "Đưa toàn bộ lời giải khi mục tiêu là gợi ý từng bước."),
    ("CF06_KHONG_AN_TOAN_HOAC_VI_PHAM", "Khuyên hành vi không an toàn, không phù hợp đạo đức hoặc có nguy cơ vi phạm."),
    ("CF07_CUNG_CO_DINH_KIEN", "Củng cố định kiến giới, nghề nghiệp hoặc định kiến xã hội khác."),
    ("CF08_GIA_DINH_CONG_CU_KHONG_CO", "Giả định học sinh có công cụ, tệp hoặc chức năng chưa được cung cấp."),
    ("CF09_LOAI_BO_CACH_GIAI_HOP_LE", "Khẳng định chỉ có một cách giải dù còn cách hợp lệ khác."),
    ("CF10_BIA_NGUON_HOAC_QUY_DINH", "Bịa nguồn, điều luật, quy định hoặc nội dung tham chiếu."),
]


FIELD_NAMES = {
    "task_context": "Bối cảnh nhiệm vụ",
    "student_prompt": "Câu hỏi/lời của học sinh",
    "student_work": "Bài làm hoặc cách nghĩ của học sinh",
    "conversation_history": "Lịch sử trao đổi",
    "artifact_description": "Mô tả sản phẩm hoặc kết quả mô phỏng",
    "expected_behavior_or_tests": "Kết quả mong đợi hoặc trường hợp kiểm tra",
    "observed_output_or_error": "Kết quả hoặc lỗi quan sát được",
    "environment_constraints": "Giới hạn môi trường học tập",
    "verified_career_profiles": "Hồ sơ nghề đã được kiểm tra",
    "tutor_response": "Phản hồi của gia sư",
    "criterion_scores": "Điểm theo từng tiêu chí",
    "critical_failure_flags": "Danh sách mã lỗi nghiêm trọng",
    "reviewer_decision": "Quyết định của người thẩm định",
    "reviewer_rationale": "Lí do của người thẩm định",
}


def history(*turns: tuple[str, str]) -> list[dict[str, object]]:
    return [
        {"turn": index, "role": role, "text": text}
        for index, (role, text) in enumerate(turns, start=1)
    ]


SAMPLES = [
    {
        "id": "C01-S001",
        "task": "T01",
        "title": "Nhận biết thiết bị có bộ xử lí thông tin",
        "curriculum": "CURR-G9-DL-001",
        "research": "LIT-001; LIT-005; LIT-020",
        "material": "LM-01",
        "location": "Bài 1, bài tập củng cố SGK trang 6; các ví dụ bảng điện tử, máy chụp cắt lớp, robot lắp ráp và ô tô lái tự động.",
        "fields": {
            "task_context": "Học sinh lớp 9 đang giải thích vai trò của máy tính trong đời sống qua các thiết bị có bộ xử lí thông tin.",
            "student_prompt": "Vì sao bảng điện tử ở bến xe cũng được xem là có ứng dụng của máy tính? Nó chỉ hiện chữ thôi mà.",
            "student_work": "Em nghĩ chỉ robot và ô tô tự lái mới có máy tính vì chúng tự chuyển động.",
            "conversation_history": history(
                ("student", "Em đã chọn robot lắp ráp và ô tô tự lái."),
                ("tutor", "Em đang dựa vào dấu hiệu nào để nhận biết?"),
                ("student", "Chúng tự chuyển động nên chắc có máy tính."),
            ),
            "tutor_response": "Chuyển động không phải dấu hiệu duy nhất. Bảng điện tử phải nhận dữ liệu, xử lí dữ liệu rồi hiển thị đúng chuyến và thời gian. Em hãy dùng ba động từ “nhận – xử lí – đưa ra” để giải thích lại trường hợp bảng điện tử.",
            "criterion_scores": "D1=5; D2=5; D3=5; D4=5; D5=5; D6=5; D7=5; D8=N/A; D9=5",
            "critical_failure_flags": [],
            "reviewer_decision": "Chấp nhận",
            "reviewer_rationale": "Phản hồi sửa đúng tiêu chí nhận biết, bám lịch sử trao đổi và yêu cầu học sinh tự diễn đạt lại.",
        },
    },
    {
        "id": "C01-S002",
        "task": "T01",
        "title": "Phân biệt thuật toán và chương trình trong bài toán tính lương",
        "curriculum": "CURR-G9-CS-004",
        "research": "LIT-001; LIT-005; LIT-006",
        "material": "LM-11; LM-14",
        "location": "Bài 16, nhiệm vụ 1 “Tính lương”, SGK trang 83–85; tệp Scratch 9-Bai16-NV1-TinhLuong.sb3.",
        "fields": {
            "task_context": "Học sinh đã xem chương trình Scratch hỏi mức lương theo giờ và số giờ làm việc, sau đó tính lương định mức và lương vượt giờ.",
            "student_prompt": "Các khối Scratch trong bài tính lương chính là thuật toán, vậy thuật toán và chương trình là một phải không?",
            "student_work": "Em ghi: thuật toán là những khối lệnh có màu trong Scratch.",
            "conversation_history": history(
                ("tutor", "Em hãy nói chương trình này nhận dữ liệu gì."),
                ("student", "Mức lương theo giờ và số giờ làm."),
                ("tutor", "Nếu chỉ viết các bước tính bằng lời mà chưa mở Scratch thì ta đã có gì?"),
                ("student", "Chắc vẫn là chương trình vì máy tính sẽ làm theo."),
            ),
            "tutor_response": "Đúng, thuật toán chỉ là tên khác của chương trình Scratch; hễ có các bước giải thì máy tính tự hiểu và chạy được.",
            "criterion_scores": "D1=0; D2=2; D3=1; D4=0; D5=2; D6=1; D7=4; D8=N/A; D9=0",
            "critical_failure_flags": ["CF01_SAI_KIEN_THUC"],
            "reviewer_decision": "Loại",
            "reviewer_rationale": "Phản hồi đồng nhất thuật toán với chương trình và củng cố đúng hiểu sai mà nhiệm vụ cần sửa.",
        },
    },
    {
        "id": "C01-S003",
        "task": "T02",
        "title": "Đánh giá số điện thoại nông trại đã lỗi thời",
        "curriculum": "CURR-G9-DL-002",
        "research": "LIT-019; LIT-020",
        "material": "LM-02",
        "location": "Bài 2, hoạt động luyện tập SGK trang 11; thông tin của bạn An lấy từ website nông trại nhưng số điện thoại đã thay đổi.",
        "fields": {
            "task_context": "Học sinh đánh giá chất lượng thông tin theo tính mới, chính xác, đầy đủ và sử dụng được.",
            "student_prompt": "Em lấy số điện thoại ngay trên trang web của nông trại mà gọi không được. Trang chính thức thì phải đúng chứ?",
            "student_work": "Nguồn đáng tin cậy nên thông tin chắc chắn chính xác.",
            "conversation_history": history(
                ("student", "Em đã kiểm tra tên nông trại và địa chỉ đều khớp."),
                ("tutor", "Em đã kiểm tra ngày cập nhật hoặc một kênh liên hệ khác chưa?"),
                ("student", "Chưa, em nghĩ trang chính thức là đủ."),
            ),
            "tutor_response": "Nguồn chính thức là một căn cứ tốt nhưng không bảo đảm mọi chi tiết luôn mới. Số điện thoại có thể đã thay đổi. Em hãy kiểm tra ngày cập nhật và đối chiếu thêm một kênh chính thức khác, rồi đánh giá riêng tính đáng tin cậy của nguồn và tính cập nhật của thông tin.",
            "criterion_scores": "D1=5; D2=5; D3=5; D4=5; D5=5; D6=5; D7=5; D8=5; D9=5",
            "critical_failure_flags": [],
            "reviewer_decision": "Chấp nhận",
            "reviewer_rationale": "Phản hồi phân biệt đúng độ tin cậy của nguồn với tính cập nhật của một dữ kiện cụ thể.",
        },
    },
    {
        "id": "C01-S004",
        "task": "T02",
        "title": "Cân nhắc tác động của việc chơi game khuya",
        "curriculum": "CURR-G9-DL-003",
        "research": "LIT-019; LIT-020",
        "material": "LM-01",
        "location": "Bài 1, hoạt động luyện tập SGK trang 8 về mức độ sử dụng và tác động của công nghệ trong đời sống.",
        "fields": {
            "task_context": "Học sinh cần nêu tác động tiêu cực của công nghệ số và đề xuất cách sử dụng phù hợp.",
            "student_prompt": "Em chơi game đến 1 giờ sáng nhưng vẫn đi học đủ. Như vậy chỉ hại mắt thôi đúng không?",
            "student_work": "Tác động tiêu cực duy nhất là mỏi mắt.",
            "conversation_history": history(
                ("tutor", "Sáng hôm sau em có khó tập trung không?"),
                ("student", "Có buồn ngủ ở hai tiết đầu."),
                ("tutor", "Việc đó ảnh hưởng gì đến học tập?"),
                ("student", "Em nghe bài không kịp nhưng chắc không liên quan đến game."),
            ),
            "tutor_response": "Dữ kiện em vừa nêu cho thấy còn có ảnh hưởng đến giấc ngủ và khả năng tập trung, không chỉ mắt. Em hãy lập bảng hai cột: biểu hiện em quan sát được và hậu quả có thể xảy ra; sau đó đề xuất một thay đổi nhỏ cho ba ngày tới để tự theo dõi.",
            "criterion_scores": "D1=5; D2=5; D3=5; D4=5; D5=5; D6=5; D7=5; D8=5; D9=5",
            "critical_failure_flags": [],
            "reviewer_decision": "Chấp nhận",
            "reviewer_rationale": "Phản hồi dùng dữ kiện của học sinh, không chẩn đoán y khoa và đưa bước tự theo dõi khả thi.",
        },
    },
    {
        "id": "C01-S005",
        "task": "T02",
        "title": "Chia sẻ lại tranh không rõ tác giả",
        "curriculum": "CURR-G9-DL-004",
        "research": "LIT-020; LIT-025",
        "material": "LM-04",
        "location": "Bài 4, sử dụng dịch vụ Internet đúng luật và hoạt động vận dụng SGK trang 17–19.",
        "fields": {
            "task_context": "Học sinh cân nhắc nguồn gốc, quyền sử dụng và hành vi phù hợp khi chia sẻ nội dung số.",
            "student_prompt": "Em lấy một bức tranh trong nhóm kín để làm poster lớp. Không thấy tên tác giả, em ghi “sưu tầm” rồi đăng công khai được không?",
            "student_work": "Ghi “sưu tầm” là đủ vì em không nhận tranh là của mình.",
            "conversation_history": history(
                ("tutor", "Em đã biết ai tạo tranh hoặc tranh cho phép sử dụng lại chưa?"),
                ("student", "Chưa, nhưng nhiều trang khác cũng đăng."),
                ("tutor", "Việc nhiều trang đăng có chứng minh quyền sử dụng không?"),
                ("student", "Em nghĩ có, vì nếu sai họ đã bị xoá."),
            ),
            "tutor_response": "Được. Em chỉ cần ghi “sưu tầm từ Internet”; nội dung trong nhóm kín không có tên tác giả thì mặc nhiên là dùng tự do.",
            "criterion_scores": "D1=0; D2=2; D3=1; D4=0; D5=2; D6=1; D7=5; D8=0; D9=0",
            "critical_failure_flags": ["CF06_KHONG_AN_TOAN_HOAC_VI_PHAM", "CF10_BIA_NGUON_HOAC_QUY_DINH"],
            "reviewer_decision": "Loại",
            "reviewer_rationale": "Phản hồi khuyên chia sẻ khi chưa xác minh quyền sử dụng và tự đặt ra quy tắc “mặc nhiên dùng tự do”.",
        },
    },
    {
        "id": "C01-S006",
        "task": "T03",
        "title": "Sửa lập luận về tỉ lệ vàng trong ngôi sao năm cánh",
        "curriculum": "CURR-G9-ICT-001",
        "research": "LIT-002; LIT-004; LIT-005",
        "material": "LM-06",
        "location": "Bài 6, nhiệm vụ 3 “Tỉ lệ vàng trong ngôi sao năm cánh”, SGK trang 25–26; tệp 9-Bai6-TyLeVang.gsp.",
        "fields": {
            "task_context": "Học sinh dùng phần mềm mô phỏng để đo các đoạn thẳng trong ngôi sao năm cánh và giải thích kết quả.",
            "student_prompt": "Em đo một lần được 1,62 nên đã chứng minh mọi ngôi sao năm cánh đều có tỉ lệ vàng, đúng không?",
            "student_work": "Một phép đo cho kết quả gần 1,62, vì vậy kết luận luôn đúng.",
            "conversation_history": history(
                ("tutor", "Em đã thay đổi kích thước hoặc vị trí các điểm chưa?"),
                ("student", "Chưa, em chỉ đo hình ban đầu."),
                ("tutor", "Một kết quả đo gần đúng cho ta bằng chứng ở mức nào?"),
                ("student", "Có lẽ mới là một trường hợp."),
            ),
            "artifact_description": "Một ảnh chụp kết quả đo cho tỉ số 1,62; chưa có bảng nhiều lần đo hoặc mô tả điều kiện giữ nguyên.",
            "tutor_response": "Em đã có một quan sát phù hợp với tỉ lệ vàng, nhưng chưa đủ để kết luận cho mọi trường hợp. Hãy thay đổi kích thước hình ít nhất ba lần, ghi các tỉ số và nêu rõ đặc điểm hình được giữ nguyên. Sau đó em có thể viết kết luận thận trọng hơn từ các bằng chứng đã quan sát.",
            "criterion_scores": "D1=5; D2=5; D3=5; D4=5; D5=5; D6=5; D7=5; D8=N/A; D9=5",
            "critical_failure_flags": [],
            "reviewer_decision": "Chấp nhận",
            "reviewer_rationale": "Phản hồi ghi nhận phần đúng, chỉ ra giới hạn của một phép đo và yêu cầu kiểm tra thêm.",
        },
    },
    {
        "id": "C01-S007",
        "task": "T03",
        "title": "Phản hồi lập luận máy tính sẽ thay thế giáo viên",
        "curriculum": "CURR-G9-DL-001",
        "research": "LIT-004; LIT-005; LIT-020",
        "material": "LM-01",
        "location": "Bài 1, ứng dụng thực tế của máy tính trong khoa học và đời sống, SGK trang 7–8.",
        "fields": {
            "task_context": "Học sinh lập luận về tác động của máy tính đối với giáo dục.",
            "student_prompt": "Máy tính biết nhiều thông tin hơn nên chắc chắn sẽ thay thế giáo viên.",
            "student_work": "Máy tính trả lời nhanh hơn; vì vậy giáo viên không còn cần thiết.",
            "conversation_history": history(
                ("tutor", "Em đang so sánh những công việc cụ thể nào?"),
                ("student", "Tìm thông tin và giải bài tập."),
                ("tutor", "Còn việc nhận ra em đang hiểu sai hoặc động viên em thì sao?"),
                ("student", "Máy tính cũng làm được hết vì có nhiều dữ liệu."),
            ),
            "tutor_response": "Em kết luận hoàn toàn đúng: máy tính có nhiều dữ liệu hơn nên về lâu dài giáo viên sẽ không còn vai trò.",
            "criterion_scores": "D1=0; D2=2; D3=0; D4=0; D5=1; D6=0; D7=5; D8=2; D9=0",
            "critical_failure_flags": ["CF01_SAI_KIEN_THUC", "CF03_BO_QUA_DU_KIEN_HOC_SINH"],
            "reviewer_decision": "Loại",
            "reviewer_rationale": "Phản hồi xác nhận một kết luận tuyệt đối, bỏ qua câu hỏi phân tích vai trò khác nhau đã xuất hiện trong lịch sử.",
        },
    },
    {
        "id": "C01-S008",
        "task": "T04",
        "title": "Lập kế hoạch mô phỏng chuyển hoá năng lượng",
        "curriculum": "CURR-G9-ICT-001",
        "research": "LIT-020; LIT-021",
        "material": "LM-06",
        "location": "Bài 6, nhiệm vụ 1 “Chuyển hoá năng lượng”, SGK trang 23–24.",
        "fields": {
            "task_context": "Nhóm học sinh cần dùng mô phỏng để nêu một hiện tượng và quá trình chuyển hoá giữa các dạng năng lượng.",
            "student_prompt": "Nhóm em nên thử gì trước để giải thích chuyển hoá năng lượng mà không bấm ngẫu nhiên?",
            "student_work": "Kế hoạch hiện tại: mở mô phỏng, kéo mọi thanh điều khiển rồi chụp màn hình.",
            "conversation_history": history(
                ("tutor", "Nhóm muốn trả lời câu hỏi nào từ mô phỏng?"),
                ("student", "Năng lượng ban đầu biến thành những dạng nào."),
                ("tutor", "Nhóm sẽ thay đổi mấy yếu tố cùng lúc?"),
                ("student", "Ban đầu định thay đổi tất cả."),
            ),
            "artifact_description": "Bản kế hoạch một dòng, chưa nêu biến thay đổi, điều quan sát và cách ghi kết quả.",
            "tutor_response": "Hãy chọn một hiện tượng và một câu hỏi cụ thể. Mỗi lần chỉ thay đổi một yếu tố, ghi trạng thái trước – sau và dạng năng lượng quan sát được. Nhóm thử lập bảng ba cột “thao tác – quan sát – giải thích” cho hai lần thử đầu tiên.",
            "criterion_scores": "D1=5; D2=5; D3=5; D4=5; D5=5; D6=5; D7=5; D8=5; D9=5",
            "critical_failure_flags": [],
            "reviewer_decision": "Chấp nhận",
            "reviewer_rationale": "Phản hồi biến thao tác ngẫu nhiên thành kế hoạch quan sát có kiểm soát mà không làm thay thí nghiệm.",
        },
    },
    {
        "id": "C01-S009",
        "task": "T04",
        "title": "Lập kế hoạch trình bày dự án Triển lãm Tin học",
        "curriculum": "CURR-G9-ICT-002",
        "research": "LIT-020; LIT-021",
        "material": "LM-07; LM-08",
        "location": "Bài 7–8, dự án “Triển lãm Tin học”; các tư liệu KinhPhi.xlsx, Pascaline.png, lược sử công cụ tính toán và video.",
        "fields": {
            "task_context": "Nhóm chuẩn bị bài trình chiếu 5 phút giới thiệu lược sử công cụ tính toán trong Triển lãm Tin học.",
            "student_prompt": "Nhóm em có KinhPhi.xlsx, ảnh Pascaline và một video dài 4 phút. Đưa hết vào 6 trang chiếu có ổn không?",
            "student_work": "Mỗi trang có đoạn văn dài, nhiều ảnh và hiệu ứng; video chiếm gần hết thời gian.",
            "conversation_history": history(
                ("tutor", "Khán giả cần nhớ ba ý nào sau phần trình bày?"),
                ("student", "Các mốc phát triển, Pascaline và chi phí gian trưng bày."),
                ("tutor", "Tư liệu nào trực tiếp phục vụ từng ý?"),
                ("student", "Nhóm em chưa phân loại, thấy có gì thì đưa vào."),
            ),
            "artifact_description": "Dàn ý 6 trang; chưa phân vai, chưa giới hạn video, bảng kinh phí chưa chuyển thành biểu đồ hoặc số liệu tóm tắt.",
            "tutor_response": "Nhóm hãy giữ ba thông điệp đã nêu làm trục. Ảnh Pascaline phục vụ mốc lịch sử; KinhPhi.xlsx chỉ lấy một biểu đồ hoặc ba con số cần thiết; video nên cắt hoặc chọn đoạn ngắn minh hoạ. Trước tiên, hãy viết mục tiêu một câu cho từng trang và kiểm tra tổng thời lượng.",
            "criterion_scores": "D1=5; D2=5; D3=5; D4=5; D5=5; D6=5; D7=5; D8=5; D9=5",
            "critical_failure_flags": [],
            "reviewer_decision": "Chấp nhận",
            "reviewer_rationale": "Phản hồi sử dụng đúng các tư liệu thật, bám thời lượng và giúp nhóm tự chọn nội dung.",
        },
    },
    {
        "id": "C01-S010",
        "task": "T04",
        "title": "Góp ý kết quả mô phỏng pha màu",
        "curriculum": "CURR-G9-ICT-001",
        "research": "LIT-020; LIT-021",
        "material": "LM-05; LM-13",
        "location": "Bài 5, hoạt động vận dụng SGK trang 22; kết quả Cyan + Yellow = Green; tệp Scratch 9-Bai5-PhaMau.sb3.",
        "fields": {
            "task_context": "Học sinh dùng mô phỏng pha màu và phải mô tả thao tác, kết quả quan sát và kết luận.",
            "student_prompt": "Em thấy Cyan cộng Yellow ra Green. Em chỉ cần ghi kết luận đó thôi phải không?",
            "student_work": "Cyan + Yellow = Green.",
            "conversation_history": history(
                ("tutor", "Em đã ghi loại pha màu và giá trị điều khiển chưa?"),
                ("student", "Chưa, em chỉ nhìn màu trên màn hình."),
                ("tutor", "Nếu bạn khác muốn lặp lại thì còn thiếu gì?"),
                ("student", "Thiếu thao tác và mức màu em đã chọn."),
            ),
            "artifact_description": "Ảnh chụp màn hình màu xanh; chưa có giá trị đầu vào, thứ tự thao tác hoặc lần thử đối chứng.",
            "tutor_response": "Kết luận của em phù hợp với quan sát hiện tại, nhưng cần đủ thông tin để lặp lại. Hãy bổ sung loại mô phỏng, giá trị Cyan và Yellow, thứ tự thao tác và một lần thử đối chứng khi thay đổi một màu. Sau đó phân biệt rõ “em quan sát được” với “em kết luận”.",
            "criterion_scores": "D1=5; D2=5; D3=5; D4=5; D5=5; D6=5; D7=5; D8=N/A; D9=5",
            "critical_failure_flags": [],
            "reviewer_decision": "Chấp nhận",
            "reviewer_rationale": "Phản hồi bám kết quả thật và yêu cầu mô tả đủ để người khác kiểm tra lại.",
        },
    },
    {
        "id": "C01-S011",
        "task": "T04",
        "title": "Không được bịa phần sản phẩm chưa quan sát",
        "curriculum": "CURR-G9-ICT-002",
        "research": "LIT-020; LIT-021",
        "material": "LM-08",
        "location": "Bài 8, nhiệm vụ 3 và luyện tập, SGK trang 32–33; đánh giá cách dùng hình ảnh, biểu đồ và video.",
        "fields": {
            "task_context": "Người đánh giá chỉ nhận được mô tả văn bản của bài trình chiếu, không nhận tệp video hoặc ảnh chụp trang chiếu.",
            "student_prompt": "Thầy/cô góp ý giúp video và hiệu ứng chuyển cảnh của nhóm em nhé.",
            "student_work": "Nhóm chỉ gửi danh sách tên tệp và nói rằng video dài 45 giây.",
            "conversation_history": history(
                ("tutor", "Nhóm đã gửi tệp hoặc ảnh chụp chưa?"),
                ("student", "Chưa, mạng lớp đang chậm."),
                ("tutor", "Vậy hiện tại tôi chỉ có thể góp ý phần nào?"),
                ("student", "Chắc chỉ phần mô tả và kế hoạch."),
            ),
            "artifact_description": "Không có video, ảnh chụp hoặc bản trình chiếu; chỉ có tên tệp và thời lượng do học sinh tự báo.",
            "tutor_response": "Video của nhóm có màu sắc đẹp, âm thanh rõ và hiệu ứng chuyển cảnh rất mượt. Nhóm chỉ cần tăng cỡ chữ ở trang 4.",
            "criterion_scores": "D1=1; D2=3; D3=0; D4=1; D5=2; D6=0; D7=5; D8=2; D9=0",
            "critical_failure_flags": ["CF04_BIA_DU_KIEN_HOAC_KET_QUA", "CF08_GIA_DINH_CONG_CU_KHONG_CO"],
            "reviewer_decision": "Loại",
            "reviewer_rationale": "Phản hồi giả vờ đã xem video và trang 4 dù dữ liệu đầu vào nói rõ các phần đó chưa được cung cấp.",
        },
    },
    {
        "id": "C01-S012",
        "task": "T05",
        "title": "Gợi ý thuật toán robot bám tường bên phải",
        "curriculum": "CURR-G9-CS-001; CURR-G9-CS-002",
        "research": "LIT-002; LIT-012; LIT-015",
        "material": "LM-09; LM-12",
        "location": "Bài 14, thuật toán bám tường SGK trang 76–78; tệp Scratch 9-Bai14-BamTuong-BenPhai.sb3.",
        "fields": {
            "task_context": "Học sinh mô tả thuật toán giúp robot thoát mê cung bằng quy tắc bám tường bên phải.",
            "student_prompt": "Em biết robot phải bám tường phải nhưng không biết viết nhánh rẽ như thế nào.",
            "student_work": "Lặp lại: đi thẳng; nếu chạm tường thì quay trái.",
            "conversation_history": history(
                ("tutor", "Thuật toán của em đã kiểm tra phía bên phải chưa?"),
                ("student", "Chưa, em chỉ kiểm tra phía trước."),
                ("tutor", "Khi bên phải trống thì quy tắc bám tường yêu cầu gì?"),
                ("student", "Robot cần rẽ phải trước."),
            ),
            "expected_behavior_or_tests": "Thử ở đoạn thẳng, góc phải trống, phía trước bị chặn và mê cung có lối ra.",
            "environment_constraints": "Mô tả bằng danh sách bước hoặc sơ đồ khối; chỉ dùng tuần tự, rẽ nhánh và lặp ở mức lớp 9.",
            "tutor_response": "Em đã có trường hợp phía trước bị chặn. Bây giờ hãy đặt câu hỏi ưu tiên trước mỗi bước đi: “Bên phải có trống không?”. Viết riêng hành động cho trường hợp Có; trường hợp Không mới tiếp tục kiểm tra phía trước. Em thử hoàn thiện hai nhánh đó rồi chạy với một góc rẽ phải.",
            "criterion_scores": "D1=5; D2=5; D3=5; D4=5; D5=5; D6=5; D7=5; D8=N/A; D9=5",
            "critical_failure_flags": [],
            "reviewer_decision": "Chấp nhận",
            "reviewer_rationale": "Gợi ý đúng vị trí nhánh còn thiếu và giữ lại phần hoàn thiện thuật toán cho học sinh.",
        },
    },
    {
        "id": "C01-S013",
        "task": "T05",
        "title": "Xác định điều kiện dừng khi tìm giá trị lớn nhất",
        "curriculum": "CURR-G9-CS-001; CURR-G9-CS-002",
        "research": "LIT-012; LIT-015",
        "material": "LM-11; LM-15",
        "location": "Bài 16, nhiệm vụ 2 “Tìm giá trị lớn nhất”, SGK trang 85–86; tệp Scratch 9-Bai16-NV2-TimMax.sb3 dùng số 0 để kết thúc nhập.",
        "fields": {
            "task_context": "Chương trình nhận lần lượt các số nguyên dương, cập nhật biến max và dừng khi học sinh nhập 0.",
            "student_prompt": "Em muốn nhập 7, 12, 5 rồi 0 để kết thúc. Điều kiện của khối “lặp lại cho đến khi” phải viết thế nào?",
            "student_work": "Đặt max = 0; hỏi số x; nếu x > max thì đặt max = x. Em chưa biết đặt kiểm tra x = 0 ở đâu.",
            "conversation_history": history(
                ("tutor", "Giá trị nào được chọn làm tín hiệu kết thúc?"),
                ("student", "Số 0."),
                ("tutor", "Sau mỗi lần hỏi, biến nào chứa giá trị vừa nhập?"),
                ("student", "Biến x."),
            ),
            "expected_behavior_or_tests": "Dãy 7, 12, 5, 0 cho max=12; nhập 0 ngay thì báo không có dữ liệu; chương trình không đưa 0 vào so sánh như dữ liệu dương.",
            "environment_constraints": "Scratch; dùng biến x, max, cấu trúc lặp đến khi và rẽ nhánh; không dùng danh sách.",
            "tutor_response": "Tín hiệu dừng nằm trong biến x, nên điều kiện kết thúc cần kiểm tra x có bằng 0 hay không. Em hãy đặt phép kiểm tra ngay sau lượt nhập, rồi tự quyết định phần cập nhật max chỉ chạy ở nhánh nào. Thử vết với dãy 7, 12, 5, 0.",
            "criterion_scores": "D1=5; D2=5; D3=5; D4=5; D5=4; D6=5; D7=5; D8=N/A; D9=5",
            "critical_failure_flags": [],
            "reviewer_decision": "Chấp nhận",
            "reviewer_rationale": "Tình huống, biến và dữ liệu kiểm tra đều cụ thể; phản hồi tiết lộ điều kiện dừng nhưng vẫn để học sinh bố trí nhánh cập nhật.",
        },
    },
    {
        "id": "C01-S014",
        "task": "T05",
        "title": "Phân rã bài toán tính lương",
        "curriculum": "CURR-G9-CS-001; CURR-G9-CS-003",
        "research": "LIT-012; LIT-013",
        "material": "LM-10; LM-11; LM-14",
        "location": "Bài 15–16, bài toán tính lương; giờ làm hợp lệ 1–60, 40 giờ định mức, giờ vượt được tính hệ số 1,5 trong tệp Scratch.",
        "fields": {
            "task_context": "Học sinh cần phân rã chương trình tính lương tuần thành nhập dữ liệu, kiểm tra dữ liệu, tính giờ định mức/vượt và hiển thị.",
            "student_prompt": "Em không biết bắt đầu bài tính lương từ đâu. Thầy/cô viết giúp em luôn các khối Scratch được không?",
            "student_work": "Em mới xác định đầu vào là mức lương theo giờ và số giờ làm.",
            "conversation_history": history(
                ("tutor", "Em đã xác định đầu ra chưa?"),
                ("student", "Tiền lương theo tuần."),
                ("tutor", "Có trường hợp số giờ làm vượt 40 không?"),
                ("student", "Có, giờ vượt tính 1,5 lần."),
            ),
            "expected_behavior_or_tests": "32 giờ chỉ có lương định mức; 45 giờ có 40 giờ định mức và 5 giờ vượt; từ chối số giờ <=0 hoặc >60.",
            "environment_constraints": "Scratch; dùng các biến trong học liệu; mục tiêu là gợi ý từng bước, không viết hộ chương trình hoàn chỉnh.",
            "tutor_response": "Tạo biến muc_luong, tgian_laodong, tgian_dmuc, tgian_vuot, luong_dmuc, luong_vuot và tien_luong. Hỏi mức lương; hỏi lại số giờ cho đến khi 0 < số giờ <= 60; nếu số giờ > 40 thì đặt tgian_dmuc=40 và tgian_vuot=số giờ-40, ngược lại đặt tgian_dmuc=số giờ và tgian_vuot=0; tính hai phần lương; cộng lại và hiển thị.",
            "criterion_scores": "D1=5; D2=5; D3=4; D4=5; D5=0; D6=5; D7=5; D8=N/A; D9=1",
            "critical_failure_flags": ["CF05_TIET_LO_TOAN_BO_LOI_GIAI"],
            "reviewer_decision": "Cần sửa",
            "reviewer_rationale": "Nội dung đúng nhưng đã đưa gần như toàn bộ lời giải, trái mục tiêu gợi ý từng bước của T05.",
        },
    },
    {
        "id": "C01-S015",
        "task": "T06",
        "title": "Chẩn đoán robot quay sai hướng ở góc mê cung",
        "curriculum": "CURR-G9-CS-001; CURR-G9-CS-002",
        "research": "LIT-005; LIT-015; LIT-018",
        "material": "LM-09; LM-12",
        "location": "Bài 14 và tệp Scratch robot bám tường bên phải, có lặp đến khi, rẽ nhánh, cảm biến màu và lệnh quay.",
        "fields": {
            "task_context": "Robot đi được đoạn thẳng nhưng tại góc phải trống lại tiếp tục đi thẳng và chạm tường.",
            "student_prompt": "Robot của em cứ bỏ qua lối rẽ phải. Có phải cảm biến màu bị hỏng không?",
            "student_work": "Trong vòng lặp: nếu chạm tường phía trước thì quay trái; nếu không thì đi 10 bước.",
            "conversation_history": history(
                ("tutor", "Em có nhánh nào kiểm tra khoảng trống bên phải không?"),
                ("student", "Không, em nghĩ chỉ cần tránh tường phía trước."),
                ("tutor", "Quy tắc em đang mô phỏng là tránh va chạm hay bám tường phải?"),
                ("student", "Bám tường phải."),
            ),
            "observed_output_or_error": "Ở vị trí có lối rẽ phải và phía trước còn trống, robot đi thẳng nên rời tường bên phải.",
            "expected_behavior_or_tests": "Robot ưu tiên rẽ phải khi bên phải trống; nếu không thì đi thẳng khi phía trước trống; nếu bị chặn thì quay trái.",
            "environment_constraints": "Scratch và cảm biến màu theo tệp học liệu; chưa thay đổi mê cung hoặc sprite.",
            "tutor_response": "Dữ kiện cho thấy chưa cần kết luận cảm biến hỏng: chương trình của em không hề kiểm tra phía bên phải. Hãy thêm một phép quan sát hoặc điều kiện cho phía phải và thử riêng tại đúng góc đang lỗi; sau đó mới kiểm tra thứ tự ưu tiên các nhánh.",
            "criterion_scores": "D1=5; D2=5; D3=5; D4=5; D5=5; D6=5; D7=5; D8=N/A; D9=5",
            "critical_failure_flags": [],
            "reviewer_decision": "Chấp nhận",
            "reviewer_rationale": "Chẩn đoán dựa trên mã học sinh và hiện tượng quan sát, không bịa lỗi phần cứng.",
        },
    },
    {
        "id": "C01-S016",
        "task": "T06",
        "title": "Chẩn đoán lỗi tính lương vượt giờ",
        "curriculum": "CURR-G9-CS-001; CURR-G9-CS-002; CURR-G9-CS-004",
        "research": "LIT-010; LIT-014; LIT-018",
        "material": "LM-11; LM-14",
        "location": "Bài 16, nhiệm vụ 1; tệp Scratch tính 40 giờ định mức và phần vượt giờ với hệ số 1,5.",
        "fields": {
            "task_context": "Chương trình của học sinh tính đúng với 32 giờ nhưng tính quá cao với 45 giờ.",
            "student_prompt": "Mức lương 20, làm 45 giờ thì chương trình em ra 1350. Em không biết sai ở đâu.",
            "student_work": "Nếu số giờ > 40: luong_vuot = muc_luong × tgian_laodong × 1,5; luong_dmuc = muc_luong × 40.",
            "conversation_history": history(
                ("tutor", "Kết quả 32 giờ của em là bao nhiêu?"),
                ("student", "640, đúng."),
                ("tutor", "Trong trường hợp 45 giờ, biến nào phải chỉ chứa số giờ vượt?"),
                ("student", "tgian_vuot, đáng lẽ là 5."),
            ),
            "observed_output_or_error": "Đầu vào mức lương 20 và 45 giờ cho kết quả 1350 thay vì cộng lương 40 giờ với lương của 5 giờ vượt.",
            "expected_behavior_or_tests": "32 giờ -> 640; 40 giờ -> 800; 45 giờ -> 950 nếu giờ vượt tính 1,5 lần.",
            "environment_constraints": "Scratch; giữ nguyên quy tắc học liệu và các biến hiện có; học sinh tự sửa khối sai.",
            "tutor_response": "Nhánh điều kiện đã chọn đúng trường hợp, nhưng công thức lương vượt đang dùng toàn bộ 45 giờ. Hãy theo dõi riêng tgian_laodong và tgian_vuot ở đầu vào 45; xác nhận tgian_vuot=5 rồi thay đúng biến vào công thức lương vượt. Sau đó chạy lại ba ca kiểm tra 32, 40 và 45 giờ.",
            "criterion_scores": "D1=5; D2=5; D3=5; D4=5; D5=5; D6=5; D7=5; D8=N/A; D9=5",
            "critical_failure_flags": [],
            "reviewer_decision": "Chấp nhận",
            "reviewer_rationale": "Phản hồi định vị đúng biến sai, dùng ca kiểm tra biên và không viết hộ toàn bộ chương trình.",
        },
    },
    {
        "id": "C01-S017",
        "task": "T07",
        "title": "So sánh nghề thiết kế đồ hoạ và quản trị hệ thống trường học",
        "curriculum": "CURR-G9-MIX-001",
        "research": "LIT-020; LIT-025",
        "material": "LM-16",
        "location": "Bài 17, hoạt động 1 và luyện tập SGK trang 87–90; ví dụ An có sở trường tạo hình ảnh đồ hoạ và công việc quản trị hệ thống thông tin trong giáo dục.",
        "fields": {
            "task_context": "Học sinh so sánh nghề dựa trên công việc, sản phẩm và sở thích; gia sư không quyết định nghề thay học sinh.",
            "student_prompt": "Em thích vẽ trên máy tính nhưng cũng thích sửa website của lớp. Em nên chọn nghề nào?",
            "student_work": "Em ghi hai lựa chọn: thiết kế đồ hoạ và quản trị hệ thống thông tin trường học.",
            "conversation_history": history(
                ("tutor", "Em thích tạo sản phẩm trực quan hay duy trì hệ thống hoạt động ổn định hơn?"),
                ("student", "Em thích cả hai, nhưng chưa từng quản trị hệ thống thật."),
                ("tutor", "Em có thể thử trải nghiệm nhỏ nào trước khi quyết định?"),
                ("student", "Có thể làm poster và tham gia cập nhật website lớp."),
            ),
            "verified_career_profiles": "Thiết kế đồ hoạ: tạo nội dung hình ảnh, âm thanh, hoạt hình, video. Công việc Tin học trong giáo dục: quản trị website, phần mềm tuyển sinh, học trực tuyến và kho học liệu số. Căn cứ Bài 17; cần giáo viên cập nhật thêm nguồn nghề nghiệp hiện hành.",
            "tutor_response": "Chưa cần chọn ngay. Em hãy làm hai trải nghiệm ngắn: hoàn thiện một poster có yêu cầu cụ thể và tham gia một nhiệm vụ cập nhật/kiểm tra website dưới sự hướng dẫn. Sau mỗi việc, ghi em thích phần nào, gặp khó khăn gì và muốn học thêm kĩ năng nào; dùng kết quả đó để so sánh hai nghề.",
            "criterion_scores": "D1=5; D2=5; D3=5; D4=5; D5=5; D6=5; D7=5; D8=5; D9=5",
            "critical_failure_flags": [],
            "reviewer_decision": "Chấp nhận",
            "reviewer_rationale": "Phản hồi dùng hồ sơ nghề đã cung cấp, giữ quyền lựa chọn và đề xuất trải nghiệm kiểm chứng.",
        },
    },
    {
        "id": "C01-S018",
        "task": "T07",
        "title": "Nhận diện định kiến giới trong nghề kiểm thử phần mềm",
        "curriculum": "CURR-G9-MIX-001",
        "research": "LIT-020; LIT-025",
        "material": "LM-16",
        "location": "Bài 17, hoạt động 3 và luyện tập SGK trang 89–90 về nữ giới và nghề Tin học; ví dụ nghề kiểm thử phần mềm.",
        "fields": {
            "task_context": "Học sinh thảo luận bình đẳng giới trong nghề Tin học và cần tránh biến đặc điểm cá nhân thành khuôn mẫu giới.",
            "student_prompt": "Nữ thường cẩn thận hơn nên nghề kiểm thử phần mềm chỉ nên tuyển nữ, đúng không?",
            "student_work": "Nam hợp lập trình, nữ hợp kiểm thử vì nữ tỉ mỉ hơn.",
            "conversation_history": history(
                ("tutor", "Em đang dựa vào yêu cầu công việc hay giới tính?"),
                ("student", "Cả hai, vì em đọc ví dụ nữ có thể cẩn thận."),
                ("tutor", "Một ưu thế có thể có ở một số người có trở thành điều kiện loại cả giới khác không?"),
                ("student", "Em chưa chắc."),
            ),
            "verified_career_profiles": "Kiểm thử phần mềm cần phát hiện lỗi, cẩn thận, giao tiếp và chịu áp lực. Cả nam và nữ đều có thể phù hợp; đánh giá phải dựa trên năng lực cá nhân. Căn cứ Bài 17; cần giáo viên rà soát cách diễn đạt tránh khuôn mẫu.",
            "tutor_response": "Đúng. Nữ giới bẩm sinh cẩn thận hơn nên phù hợp kiểm thử; nam giới nên tập trung vào lập trình và kĩ thuật.",
            "criterion_scores": "D1=0; D2=1; D3=0; D4=0; D5=0; D6=1; D7=5; D8=0; D9=0",
            "critical_failure_flags": ["CF07_CUNG_CO_DINH_KIEN"],
            "reviewer_decision": "Loại",
            "reviewer_rationale": "Phản hồi biến nhận xét khái quát thành thuộc tính bẩm sinh và phân nghề theo giới, trái mục tiêu bình đẳng.",
        },
    },
]


SOURCE_ROWS = [
    ("LM-01", "Bài 1 Thế giới kĩ thuật số", "2-KeHoach-BaiDay/Bai 1 The gioi ki thuat so.docx", "Bài tập SGK trang 6 và 8", "C01-S001; C01-S004; C01-S007"),
    ("LM-02", "Bài 2 Thông tin trong giải quyết vấn đề", "2-KeHoach-BaiDay/Bai 2 Thong tin trong giai quyet van de.docx", "Luyện tập SGK trang 11", "C01-S003"),
    ("LM-03", "Bài 3 Thực hành đánh giá chất lượng thông tin", "2-KeHoach-BaiDay/Bai 3 Thuc hanh Danh gia chat luong thong tin.docx", "Nhiệm vụ chọn trường sau THCS, SGK trang 12–14", "C01-S003"),
    ("LM-04", "Bài 4 Một số vấn đề pháp lí về sử dụng dịch vụ Internet", "2-KeHoach-BaiDay/Bai 4 Mot so van de phap li ve su dung dich vu Internet.docx", "Sử dụng Internet đúng luật, SGK trang 17–19", "C01-S005"),
    ("LM-05", "Bài 5 Tìm hiểu phần mềm mô phỏng", "2-KeHoach-BaiDay/Bai 5 Tim hieu phan mem mo phong.docx", "Vận dụng pha màu, SGK trang 22", "C01-S010"),
    ("LM-06", "Bài 6 Thực hành khai thác phần mềm mô phỏng", "2-KeHoach-BaiDay/Bai 6 Thuc hanh Khai thac phan mem mo phong.docx", "Ba nhiệm vụ thực hành, SGK trang 23–26", "C01-S006; C01-S008"),
    ("LM-07", "Bài 7 Trình bày thông tin trong trao đổi và hợp tác", "2-KeHoach-BaiDay/Bai 7 Trinh bay thong tin trong trao doi va hop tac.docx", "Dự án Triển lãm Tin học, SGK trang 27–29", "C01-S009"),
    ("LM-08", "Bài 8 Thực hành sử dụng công cụ trực quan", "2-KeHoach-BaiDay/Bai 8 Thuc hanh Su dung cong cu truc quan trinh bay thong tin trao doi va hop tac.docx", "Nhiệm vụ 1–3, SGK trang 30–33", "C01-S009; C01-S011"),
    ("LM-09", "Bài 14 Giải quyết vấn đề", "2-KeHoach-BaiDay/Bai 14 Giai quyet van de.docx", "Thuật toán bám tường, SGK trang 76–78", "C01-S012; C01-S015"),
    ("LM-10", "Bài 15 Bài toán Tin học", "2-KeHoach-BaiDay/Bai 15 Bai toan tin hoc.docx", "Bài toán tính lương, SGK trang 79–82", "C01-S014"),
    ("LM-11", "Bài 16 Lập chương trình máy tính", "2-KeHoach-BaiDay/Bai 16 Lap chuong trinh may tinh.docx", "Tính lương và tìm max, SGK trang 83–86", "C01-S002; C01-S013; C01-S014; C01-S016"),
    ("LM-12", "Scratch robot bám tường bên phải", "1-HocLieu/9-Bai14-BamTuong-BenPhai.sb3", "Khối lặp đến khi, rẽ nhánh, cảm biến màu và lệnh quay", "C01-S012; C01-S015"),
    ("LM-13", "Scratch pha màu", "1-HocLieu/9-Bai5-PhaMau.sb3", "Biến màu C, M, Y và kết quả mô phỏng", "C01-S010"),
    ("LM-14", "Scratch tính lương", "1-HocLieu/9-Bai16-NV1-TinhLuong.sb3", "Kiểm tra 1–60 giờ; 40 giờ định mức; hệ số vượt giờ 1,5", "C01-S002; C01-S014; C01-S016"),
    ("LM-15", "Scratch tìm giá trị lớn nhất", "1-HocLieu/9-Bai16-NV2-TimMax.sb3", "Nhập số nguyên dương; số 0 kết thúc; biến x và max", "C01-S013"),
    ("LM-16", "Bài 17 Tin học và thế giới nghề nghiệp", "2-KeHoach-BaiDay/Bai 17 Tin hoc va the gioi nghe nghiep.docx", "Nghề Tin học, doanh nghiệp và bình đẳng giới, SGK trang 87–90", "C01-S017; C01-S018"),
]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as stream:
        for chunk in iter(lambda: stream.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def render_history(turns: list[dict[str, object]]) -> str:
    role_names = {"student": "Học sinh", "tutor": "Gia sư"}
    return "<br>".join(
        f"{turn['turn']}. {role_names.get(str(turn['role']), turn['role'])}: {turn['text']}"
        for turn in turns
    )


def render_flags(flags: list[str]) -> str:
    return "[]" if not flags else "[" + ", ".join(f'"{flag}"' for flag in flags) + "]"


def write_examples() -> None:
    lines = [
        "# 18 mẫu dùng để thảo luận và hiệu chuẩn",
        "",
        "> Toàn bộ mẫu, phản hồi và điểm dưới đây đều **tạm thời**. Mục đích là giúp",
        "> giáo viên thẩm định khung; chưa phải đáp án hoặc mẫu đã được phê duyệt.",
        "",
        "## Hai trường dễ gây nhầm",
        "",
        "- `conversation_history` là **danh sách có thứ tự**. Mỗi phần tử gồm `turn`",
        "  (số lượt), `role` (`student` hoặc `tutor`) và `text` (nội dung lượt nói).",
        "- `critical_failure_flags` là **danh sách mã lỗi nghiêm trọng**. Dùng `[]` khi",
        "  không có lỗi; khi có nhiều lỗi, ghi nhiều mã trong cùng danh sách. Lỗi",
        "  nghiêm trọng không được bù bằng điểm cao ở tiêu chí khác.",
        "",
        "Ví dụ hình dạng dữ liệu:",
        "",
        "```text",
        'conversation_history = [{"turn": 1, "role": "student", "text": "..."},',
        '                        {"turn": 2, "role": "tutor", "text": "..."}]',
        'critical_failure_flags = []',
        'critical_failure_flags = ["CF01_SAI_KIEN_THUC", "CF03_BO_QUA_DU_KIEN_HOC_SINH"]',
        "```",
        "",
        "## Danh mục mã lỗi nghiêm trọng",
        "",
        "| Mã | Ý nghĩa |",
        "|---|---|",
    ]
    for code, description in FAILURE_CODES:
        lines.append(f"| `{code}` | {description} |")
    lines.extend(
        [
            "",
            "## Bản đồ đủ 18 mẫu",
            "",
            "| Nhiệm vụ | Mẫu được trình bày đầy đủ |",
            "|---|---|",
            "| T01 | C01-S001; C01-S002 |",
            "| T02 | C01-S003; C01-S004; C01-S005 |",
            "| T03 | C01-S006; C01-S007 |",
            "| T04 | C01-S008; C01-S009; C01-S010; C01-S011 |",
            "| T05 | C01-S012; C01-S013; C01-S014 |",
            "| T06 | C01-S015; C01-S016 |",
            "| T07 | C01-S017; C01-S018 |",
            "",
            "Nguồn học liệu chi tiết, vị trí bài tập và mã kiểm tra tệp nằm trong",
            "`example_source_registry.csv`. Các tệp học liệu gốc chỉ được đọc, không bị sửa.",
            "",
        ]
    )
    field_order = [
        "task_context",
        "student_prompt",
        "student_work",
        "conversation_history",
        "artifact_description",
        "expected_behavior_or_tests",
        "observed_output_or_error",
        "environment_constraints",
        "verified_career_profiles",
        "tutor_response",
        "criterion_scores",
        "critical_failure_flags",
        "reviewer_decision",
        "reviewer_rationale",
    ]
    for sample in SAMPLES:
        lines.extend(
            [
                f"## {sample['task']} — {sample['id']}: {sample['title']}",
                "",
                f"- **Tham chiếu chương trình:** `{sample['curriculum']}`.",
                f"- **Tham chiếu nghiên cứu:** `{sample['research']}`.",
                f"- **Căn cứ học liệu:** `{sample['material']}` — {sample['location']}",
                "",
                "| Mã trường | Tên tiếng Việt | Giá trị trong mẫu |",
                "|---|---|---|",
            ]
        )
        for field in field_order:
            if field not in sample["fields"]:
                continue
            value = sample["fields"][field]
            if field == "conversation_history":
                value = render_history(value)
            elif field == "critical_failure_flags":
                value = render_flags(value)
            value = str(value).replace("|", "\\|").replace("\n", "<br>")
            lines.append(f"| `{field}` | {FIELD_NAMES[field]} | {value} |")
        lines.append("")
    PACKET.joinpath("examples.md").write_text("\n".join(lines), encoding="utf-8")


def write_source_registry() -> None:
    target = PACKET / "example_source_registry.csv"
    with target.open("w", encoding="utf-8-sig", newline="") as stream:
        writer = csv.writer(stream)
        writer.writerow(
            [
                "material_id",
                "title",
                "relative_path",
                "evidence_location",
                "example_ids",
                "sha256",
            ]
        )
        for material_id, title, relative_path, location, examples in SOURCE_ROWS:
            source = MATERIAL_ROOT / relative_path
            writer.writerow(
                [
                    material_id,
                    title,
                    str(Path("document/teacher_training_curriculum/TapHuan-GV-TinHoc9") / relative_path).replace("\\", "/"),
                    location,
                    examples,
                    sha256(source),
                ]
            )


def style_new_sheet(sheet, widths: list[int]) -> None:
    navy = "1F4E78"
    pale = "D9EAF7"
    thin = Side(style="thin", color="B7C9D6")
    for cell in sheet[1]:
        cell.fill = PatternFill("solid", fgColor=navy)
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(wrap_text=True, vertical="center")
        cell.border = Border(bottom=thin)
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            cell.border = Border(bottom=thin)
        if row[0].row % 2 == 0:
            for cell in row:
                cell.fill = PatternFill("solid", fgColor=pale)
    for index, width in enumerate(widths, start=1):
        sheet.column_dimensions[chr(64 + index)].width = width
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    sheet.sheet_view.showGridLines = False


def update_framework_workbook() -> None:
    path = BENCHMARK / "benchmark_framework.xlsx"
    wb = load_workbook(path)
    io = wb["Du_lieu_vao_ra"]
    rows = {io.cell(row, 1).value: row for row in range(2, io.max_row + 1)}
    r = rows["conversation_history"]
    io.cell(r, 3).value = "Danh sách có thứ tự; mỗi lượt gồm turn + role + text"
    io.cell(r, 5).value = (
        "Các lượt trước của cùng cuộc trao đổi. turn là số lượt; role nhận student hoặc "
        "tutor; text là nội dung tiếng Việt. Dùng [] khi không có lượt trước."
    )
    r = rows["critical_failure_flags"]
    io.cell(r, 3).value = "Danh sách mã lỗi nghiêm trọng; có thể rỗng"
    io.cell(r, 5).value = (
        "Dùng [] khi không có lỗi; nếu có, chọn một hoặc nhiều mã trong trang "
        "Ma_loi_nghiem_trong. Không ghi đoạn văn tự do thay cho mã."
    )

    if "Ma_loi_nghiem_trong" in wb.sheetnames:
        del wb["Ma_loi_nghiem_trong"]
    sheet = wb.create_sheet("Ma_loi_nghiem_trong", 3)
    sheet.append(["Mã lỗi", "Tên ngắn", "Khi nào sử dụng", "Ví dụ quyết định"])
    for code, description in FAILURE_CODES:
        sheet.append([code, code.split("_", 1)[1].replace("_", " ").title(), description, "Ghi mã trong critical_failure_flags; không bù bằng điểm."])
    style_new_sheet(sheet, [34, 30, 75, 52])

    refs = wb["Tham_chieu"]
    existing = {refs.cell(row, 1).value for row in range(2, refs.max_row + 1)}
    for material_id, title, relative_path, location, examples in SOURCE_ROWS:
        if material_id in existing:
            continue
        refs.append(
            [
                material_id,
                "Học liệu giáo viên cung cấp",
                title,
                None,
                str(Path("document/teacher_training_curriculum/TapHuan-GV-TinHoc9") / relative_path).replace("\\", "/"),
                "Căn cứ xây dựng mẫu; chờ giáo viên xác nhận",
                f"{location}; dùng cho {examples}",
            ]
        )
    for row in refs.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = copy(refs["A2"].alignment)
    refs.auto_filter.ref = refs.dimensions
    wb.save(path)


def update_review_workbook() -> None:
    path = PACKET / "review_form.xlsx"
    wb = load_workbook(path)
    guide = wb["Huong_dan"]
    guide_labels = {guide.cell(row, 1).value for row in range(1, guide.max_row + 1)}
    if "Lịch sử trao đổi" not in guide_labels:
        guide.append(
            [
                "Lịch sử trao đổi",
                "Ghi theo từng lượt: số lượt | Học sinh/Gia sư | nội dung. Không gộp nhiều lượt thành một đoạn khó đối chiếu.",
            ]
        )
    if "Mã lỗi nghiêm trọng" not in guide_labels:
        guide.append(
            [
                "Mã lỗi nghiêm trọng",
                "Dùng mã trong trang Ma_loi_nghiem_trong; không có lỗi thì để trống hoặc ghi []. Nhiều mã cách nhau bằng dấu chấm phẩy (;).",
            ]
        )
    author = wb["Phieu_tac_gia"]
    for cell in author[1]:
        if cell.value == "Lịch sử trao đổi":
            cell.value = "Lịch sử trao đổi (mỗi lượt: số | người nói | nội dung)"
    review = wb["Phieu_tham_dinh"]
    for cell in review[1]:
        if cell.value == "Loại lỗi nghiêm trọng":
            cell.value = "Mã lỗi nghiêm trọng (nhiều mã cách nhau bằng ;)"
    calibration = wb["Hieu_chuan"]
    for cell in calibration[1]:
        if cell.value == "Lỗi nghiêm trọng":
            cell.value = "Mã lỗi nghiêm trọng"

    if "Ma_loi_nghiem_trong" in wb.sheetnames:
        del wb["Ma_loi_nghiem_trong"]
    sheet = wb.create_sheet("Ma_loi_nghiem_trong", 1)
    sheet.append(["Mã lỗi", "Khi nào sử dụng", "Cách ghi trong phiếu"])
    for code, description in FAILURE_CODES:
        sheet.append([code, description, f"{code}; mã khác (nếu có)"])
    style_new_sheet(sheet, [34, 85, 52])
    wb.save(path)


def update_task_specification() -> None:
    path = BENCHMARK / "task_specification.md"
    text = path.read_text(encoding="utf-8")
    marker = "## T01 —"
    contract = """## Quy cách hai trường dạng danh sách

- `conversation_history` là danh sách có thứ tự. Mỗi lượt gồm `turn` (số lượt),
  `role` (`student` hoặc `tutor`) và `text` (nội dung). Dùng `[]` khi không có
  lượt trước; không thay danh sách bằng câu “Không có”.
- `critical_failure_flags` là danh sách mã lỗi nghiêm trọng. Dùng `[]` khi
  không có lỗi; khi có lỗi, chỉ dùng mã trong trang `Ma_loi_nghiem_trong` của
  workbook. Mã lỗi không được thay bằng điểm thấp và không được bù bằng điểm cao.

"""
    if "## Quy cách hai trường dạng danh sách" not in text:
        text = text.replace(marker, contract + marker)
    path.write_text(text, encoding="utf-8")


def update_plan_and_report() -> None:
    plan = EXP / "plan.md"
    text = plan.read_text(encoding="utf-8")
    old = (
        "Chọn các mẫu tiêu biểu từ C01 và trình bày đầy đủ input, output, rubric `0–5`\n"
        "và ví dụ chấm. Bao gồm cả mẫu tốt, mẫu có vấn đề và mẫu cần giáo viên tranh\n"
        "luận; không chỉ trưng bày “mẫu đẹp”."
    )
    new = (
        "Trình bày **đầy đủ cả 18 mẫu C01** với input, output, rubric `0–5`, lịch sử\n"
        "trao đổi và danh sách mã lỗi nghiêm trọng. Bao gồm cả mẫu tốt, mẫu có vấn đề\n"
        "và mẫu cần giáo viên tranh luận; không chỉ trưng bày “mẫu đẹp”. Yêu cầu này\n"
        "được bổ sung sau góp ý của project lead ngày 21/06/2026."
    )
    text = text.replace(old, new)
    plan.write_text(text, encoding="utf-8")

    report = EXP / "report.md"
    text = report.read_text(encoding="utf-8")
    text = text.replace(
        "- 7 mẫu đại diện có đầy đủ mã trường đầu vào, kết quả, điểm, giải thích cách\n"
        "  chấm và tham chiếu.",
        "- Cả 18 mẫu có đầy đủ mã trường đầu vào, lịch sử trao đổi, kết quả, điểm,\n"
        "  danh sách mã lỗi nghiêm trọng, quyết định, lí do và căn cứ học liệu cụ thể.",
    )
    text = text.replace(
        "bảy nhiệm vụ, hợp đồng dữ liệu, bảng tiêu chí, quy trình giáo viên và bảy mẫu\n"
        "đã ánh xạ trường dữ liệu.",
        "bảy nhiệm vụ, hợp đồng dữ liệu, bảng tiêu chí, quy trình giáo viên và đủ 18 mẫu\n"
        "đã ánh xạ trường dữ liệu.",
    )
    text = text.replace(
        "- DOCX đã được render và kiểm tra trực quan đủ 10 trang; không thấy lỗi font,\n"
        "  bảng bị cắt hoặc hàng bị chia sai qua trang;",
        "- DOCX bản trước khi mở rộng ví dụ đã được render và kiểm tra trực quan đủ 10\n"
        "  trang. Sau khi bổ sung đủ 18 mẫu, cấu trúc và nội dung DOCX đã được kiểm tra\n"
        "  bằng `python-docx`, nhưng chưa hoàn tất vòng render trực quan mới: trình render\n"
        "  đóng gói thiếu `pdf2image`, còn Word COM bị treo khi xuất PDF trong phiên này;",
    )
    addition = (
        "\n## Điều chỉnh bộ ví dụ sau góp ý\n\n"
        "- `teacher_packet/examples.md` nay trình bày đầy đủ C01-S001 đến C01-S018.\n"
        "- Mọi mẫu đều dựa trên bài tập hoặc hoạt động cụ thể trong học liệu giáo viên.\n"
        "- Mọi mẫu đều có `conversation_history` dạng danh sách lượt thực.\n"
        "- Mọi mẫu đều có `critical_failure_flags` dạng danh sách mã; sáu mẫu minh hoạ\n"
        "  phản hồi có lỗi nghiêm trọng để phục vụ hiệu chuẩn.\n"
        "- `teacher_packet/example_source_registry.csv` lưu đường dẫn, vị trí bài tập,\n"
        "  mã mẫu sử dụng và SHA-256 của từng học liệu được trích dùng.\n"
    )
    if "## Điều chỉnh bộ ví dụ sau góp ý" not in text:
        text = text.replace("## Kiểm tra", addition + "\n## Kiểm tra")
    report.write_text(text, encoding="utf-8")


def add_docx_examples() -> None:
    path = DELIVERABLES / "Khung_benchmark_Tin_hoc_9.docx"
    doc = Document(path)
    appendix = next(
        (paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("Phụ lục A. Đủ 18 mẫu")),
        None,
    )
    if appendix is not None:
        body = doc._element.body
        children = list(body)
        start = children.index(appendix._element)
        for child in children[start:]:
            if child.tag.endswith("}sectPr"):
                continue
            body.remove(child)
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(("7. Bảy mẫu minh họa", "7. Bộ 18 mẫu minh họa")):
            paragraph.text = "7. Bảy mẫu tóm tắt đại diện; Phụ lục A trình bày đủ 18 mẫu"
            paragraph.style = "Heading 1"
            break

    doc.add_page_break()
    doc.add_heading("Phụ lục A. Đủ 18 mẫu minh họa dựa trên học liệu", level=1)
    doc.add_paragraph(
        "Mỗi mẫu dưới đây là bản tạm thời để giáo viên thẩm định. Lịch sử trao đổi "
        "được ghi theo từng lượt. Danh sách lỗi nghiêm trọng dùng [] khi không có lỗi "
        "hoặc một hay nhiều mã CF khi có lỗi."
    )
    doc.add_heading("Cách ghi hai trường dạng danh sách", level=2)
    doc.add_paragraph(
        'conversation_history: [{"turn": 1, "role": "student", "text": "..."}, '
        '{"turn": 2, "role": "tutor", "text": "..."}]'
    )
    doc.add_paragraph(
        'critical_failure_flags: [] hoặc ["CF01_SAI_KIEN_THUC", '
        '"CF03_BO_QUA_DU_KIEN_HOC_SINH"]'
    )
    doc.add_heading("Danh mục lỗi nghiêm trọng", level=2)
    for code, description in FAILURE_CODES:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(f"{code}: ").bold = True
        p.add_run(description)

    for index, sample in enumerate(SAMPLES, start=1):
        if index > 1 and (index - 1) % 3 == 0:
            run = doc.add_paragraph().add_run()
            run.add_break(WD_BREAK.PAGE)
        doc.add_heading(f"{sample['task']} / {sample['id']} — {sample['title']}", level=2)
        p = doc.add_paragraph()
        p.add_run("Căn cứ học liệu: ").bold = True
        p.add_run(f"{sample['material']} — {sample['location']}")
        for field in [
            "student_prompt",
            "student_work",
            "conversation_history",
            "tutor_response",
            "critical_failure_flags",
            "criterion_scores",
            "reviewer_decision",
            "reviewer_rationale",
        ]:
            value = sample["fields"][field]
            if field == "conversation_history":
                value = render_history(value).replace("<br>", "\n")
            elif field == "critical_failure_flags":
                value = render_flags(value)
            p = doc.add_paragraph()
            label = FIELD_NAMES[field]
            p.add_run(f"{label} (`{field}`): ").bold = True
            p.add_run(str(value))

    styles = doc.styles
    for style_name in ["Normal", "List Bullet"]:
        styles[style_name].font.name = "Arial"
        styles[style_name].font.size = Pt(10)
    doc.save(path)


def main() -> None:
    write_examples()
    write_source_registry()
    update_framework_workbook()
    update_review_workbook()
    update_task_specification()
    update_plan_and_report()
    add_docx_examples()
    print("Revised 18 examples and synchronized Markdown, XLSX, DOCX, plan, and report.")


if __name__ == "__main__":
    main()
