from __future__ import annotations

from pathlib import Path
import re

from docx import Document


ROOT = Path("document/teacher_training_curriculum/TapHuan-GV-TinHoc9")
OUT = Path("experiments/20260621_135515/_learning_material_extract")
OUT.mkdir(parents=True, exist_ok=True)

FILES = [
    ROOT / "2-KeHoach-BaiDay/Bai 1 The gioi ki thuat so.docx",
    ROOT / "2-KeHoach-BaiDay/Bai 2 Thong tin trong giai quyet van de.docx",
    ROOT / "2-KeHoach-BaiDay/Bai 3 Thuc hanh Danh gia chat luong thong tin.docx",
    ROOT / "2-KeHoach-BaiDay/Bai 4 Mot so van de phap li ve su dung dich vu Internet.docx",
    ROOT / "2-KeHoach-BaiDay/Bai 5 Tim hieu phan mem mo phong.docx",
    ROOT / "2-KeHoach-BaiDay/Bai 6 Thuc hanh Khai thac phan mem mo phong.docx",
    ROOT / "2-KeHoach-BaiDay/Bai 7 Trinh bay thong tin trong trao doi va hop tac.docx",
    ROOT / "2-KeHoach-BaiDay/Bai 8 Thuc hanh Su dung cong cu truc quan trinh bay thong tin trao doi va hop tac.docx",
    ROOT / "2-KeHoach-BaiDay/Bai 14 Giai quyet van de.docx",
    ROOT / "2-KeHoach-BaiDay/Bai 15 Bai toan tin hoc.docx",
    ROOT / "2-KeHoach-BaiDay/Bai 16 Lap chuong trinh may tinh.docx",
    ROOT / "2-KeHoach-BaiDay/Bai 17 Tin hoc va the gioi nghe nghiep.docx",
    ROOT / "4-KiemTra-DanhGia/c-DacTa-DeKiemTra-HK1-TinHoc9.docx",
    ROOT / "5-TuLieu-ThamKhao/10-CauHoi-TracNghiem-TinHoc9.docx",
    ROOT / "5-TuLieu-ThamKhao/ViDu-BaiTap-DuAn-KeHoachBaiDay.docx",
]


def clean(text: str) -> str:
    return re.sub(r"\s+", " ", text).strip()


def extract(path: Path) -> str:
    doc = Document(path)
    lines = [f"# {path.name}", "", f"Nguồn: `{path.as_posix()}`", ""]
    for i, paragraph in enumerate(doc.paragraphs, 1):
        text = clean(paragraph.text)
        if text:
            lines.append(f"P{i}: {text}")
    for ti, table in enumerate(doc.tables, 1):
        lines.extend(["", f"## Bảng {ti}", ""])
        for ri, row in enumerate(table.rows, 1):
            values = [clean(cell.text) for cell in row.cells]
            if any(values):
                lines.append(f"R{ri}: " + " | ".join(values))
    return "\n".join(lines) + "\n"


index = ["# Trích xuất học liệu phục vụ thiết kế ví dụ", ""]
for path in FILES:
    if not path.exists():
        index.append(f"- THIẾU: `{path}`")
        continue
    target = OUT / f"{path.stem}.md"
    target.write_text(extract(path), encoding="utf-8")
    index.append(f"- [{path.name}]({target.name})")

(OUT / "README.md").write_text("\n".join(index) + "\n", encoding="utf-8")
print(f"Extracted {len(FILES)} files to {OUT}")
